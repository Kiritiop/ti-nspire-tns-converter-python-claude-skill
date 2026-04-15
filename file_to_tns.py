#!/usr/bin/env python3
"""
file_to_tns.py — Convert any document to a TI-Nspire .tns Notes file.

Supported: .pdf  .txt  .md  .docx  .xlsx  .csv  .html  .htm  .rtf

Usage:
    python3 file_to_tns.py input.pdf
    python3 file_to_tns.py notes.docx output.tns
    python3 file_to_tns.py data.csv --preview

Requirements:
    pip install pdfplumber pypdf python-docx openpyxl

Luna (auto-built on first run, requires make + zlib):
    macOS:  brew install zlib
    Linux:  sudo apt install zlib1g-dev
"""

import sys, os, re, csv, textwrap, tempfile, shutil, subprocess
import argparse, zipfile, urllib.request
from xml.sax.saxutils import escape as xml_escape

# ── Luna setup ────────────────────────────────────────────────

LUNA_BUILD_DIR = os.path.expanduser("~/.luna_build")
LUNA_BIN       = os.path.join(LUNA_BUILD_DIR, "luna")
LUNA_SEARCH    = [
    LUNA_BIN,
    os.path.expanduser("~/Desktop/Coding/Luna/luna"),
    os.path.expanduser("~/Luna/luna"),
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "luna"),
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "Luna", "luna"),
    shutil.which("luna") or "",
]
LUNA_REPO = "https://github.com/ndless-nspire/Luna/archive/refs/heads/master.zip"

_luna_path = None  # cached after first find

def find_luna():
    global _luna_path
    if _luna_path:
        return _luna_path
    for p in LUNA_SEARCH:
        if p and os.path.isfile(p) and os.access(p, os.X_OK):
            _luna_path = p
            return p
    return _build_luna()

def _build_luna():
    global _luna_path
    print("[setup] Luna not found — downloading and building …")
    os.makedirs(LUNA_BUILD_DIR, exist_ok=True)
    zip_path = os.path.join(LUNA_BUILD_DIR, "luna_src.zip")
    try:
        try:
            import ssl, certifi
            ctx = ssl.create_default_context(cafile=certifi.where())
            with urllib.request.urlopen(LUNA_REPO, context=ctx) as r:
                open(zip_path, "wb").write(r.read())
        except Exception:
            # fallback without certifi
            urllib.request.urlretrieve(LUNA_REPO, zip_path)
    except Exception as e:
        sys.exit(
            f"[ERROR] Could not download Luna: {e}\n"
            "  Manual fix: git clone https://github.com/ndless-nspire/Luna\n"
            "              cd Luna && make\n"
            f"             cp luna {LUNA_BIN}"
        )
    src_dir = os.path.join(LUNA_BUILD_DIR, "src")
    os.makedirs(src_dir, exist_ok=True)
    with zipfile.ZipFile(zip_path) as z:
        z.extractall(src_dir)
    luna_src = os.path.join(src_dir, "Luna-master")
    if not os.path.isdir(luna_src):
        for e in os.listdir(src_dir):
            if "luna" in e.lower() and os.path.isdir(os.path.join(src_dir, e)):
                luna_src = os.path.join(src_dir, e)
                break
    print("[setup] Building Luna …")
    r = subprocess.run(["make"], cwd=luna_src, capture_output=True, text=True)
    if r.returncode != 0:
        print(r.stderr)
        platform_hint = (
            "  brew install zlib" if sys.platform == "darwin"
            else "  sudo apt install zlib1g-dev"
        )
        sys.exit(f"[ERROR] Build failed. Try:\n{platform_hint}\nThen run again.")
    built = os.path.join(luna_src, "luna")
    if not os.path.isfile(built):
        sys.exit("[ERROR] Build succeeded but luna binary not found.")
    shutil.copy2(built, LUNA_BIN)
    print(f"[setup] Luna ready → {LUNA_BIN}")
    _luna_path = LUNA_BIN
    return LUNA_BIN

# ── Text extraction ───────────────────────────────────────────

def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext in (".txt", ".md"):
        return open(path, encoding="utf-8", errors="replace").read()

    if ext == ".pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        pages.append(f"=== Page {i+1} ===\n{t.strip()}")
            if pages:
                return "\n\n".join(pages)
        except ImportError:
            pass
        from pypdf import PdfReader
        pages = []
        for i, page in enumerate(PdfReader(path).pages):
            t = page.extract_text()
            if t and t.strip():
                pages.append(f"=== Page {i+1} ===\n{t.strip()}")
        return "\n\n".join(pages)

    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        # Also grab table cells
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    if ext in (".xlsx", ".xls"):
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"=== {ws.title} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    if ext == ".csv":
        lines = []
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            for row in csv.reader(f):
                lines.append(" | ".join(row))
        return "\n".join(lines)

    if ext in (".html", ".htm"):
        from html.parser import HTMLParser
        class Stripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts = []
            def handle_data(self, d):
                if d.strip():
                    self.parts.append(d.strip())
        s = Stripper()
        s.feed(open(path, encoding="utf-8", errors="replace").read())
        return "\n".join(s.parts)

    if ext == ".rtf":
        raw = open(path, encoding="utf-8", errors="replace").read()
        raw = re.sub(r'\\[a-z]+\d* ?', ' ', raw)
        raw = re.sub(r'[{}\\]', '', raw)
        return re.sub(r'  +', ' ', raw).strip()

    # Fallback — try plain text
    return open(path, encoding="utf-8", errors="replace").read()

# ── Chunking ──────────────────────────────────────────────────

MAX_CHARS = 11500

def split_text(text: str) -> list:
    if len(text) <= MAX_CHARS:
        return [text]
    paragraphs = re.split(r'\n{2,}', text)
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= MAX_CHARS:
            current += ("\n\n" if current else "") + para
        else:
            if current:
                chunks.append(current)
            if len(para) > MAX_CHARS:
                for i in range(0, len(para), MAX_CHARS):
                    chunks.append(para[i:i + MAX_CHARS])
                current = ""
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks

# ── TNS XML + Luna ────────────────────────────────────────────

WRAP = 48

def _fmtxt(text: str) -> str:
    paras = []
    for line in text.splitlines():
        if not line.strip():
            paras.append(
                '<node name="1para"><node name="1rtline">'
                '<leaf name="1word"> </leaf></node></node>')
            continue
        for wrapped in textwrap.wrap(line, width=WRAP) or [line]:
            leaves = "".join(
                f'<leaf name="1word">{xml_escape(w)} </leaf>'
                for w in wrapped.split()
            )
            paras.append(
                f'<node name="1para"><node name="1rtline">'
                f'{leaves}</node></node>')
    return xml_escape(
        '<r2dtotree><node name="1doc">'
        + "".join(paras)
        + '</node></r2dtotree>'
    )

def _problem_xml(text: str, name: str = "") -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" ?>\n'
        f'<prob xmlns="urn:TI.Problem" ver="1.0" pbname="{xml_escape(name)}">'
        '<sym></sym>'
        '<card clay="0" h1="10000" h2="10000" w1="10000" w2="10000">'
        '<isDummyCard>0</isDummyCard><flag>0</flag>'
        '<wdgt xmlns:np="urn:TI.Notepad" type="TI.Notepad" ver="2.0">'
        '<np:mFlags>1024</np:mFlags>'
        f'<np:value>{len(text)}</np:value>'
        f'<np:fmtxt>{_fmtxt(text)}</np:fmtxt>'
        '</wdgt></card></prob>'
    )

def build_tns(text: str, out_path: str, doc_name: str = "notes") -> int:
    luna = find_luna()
    chunks = split_text(text)
    with tempfile.TemporaryDirectory() as tmp:
        xmls = []
        for i, chunk in enumerate(chunks):
            name = f"p{i+1}" if len(chunks) > 1 else doc_name[:20]
            xml_path = os.path.join(tmp, f"Problem{i+1}.xml")
            open(xml_path, "w", encoding="utf-8").write(_problem_xml(chunk, name))
            xmls.append(xml_path)
        r = subprocess.run([luna] + xmls + [out_path],
                           capture_output=True, text=True)
        if r.returncode != 0:
            sys.exit(f"[ERROR] Luna failed:\n{r.stdout}\n{r.stderr}")
    return len(chunks)

# ── CLI ───────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Convert any document to a TI-Nspire .tns Notes file.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            Supported formats: .pdf .txt .md .docx .xlsx .csv .html .htm .rtf

            Examples:
              python3 file_to_tns.py notes.pdf
              python3 file_to_tns.py chem.docx chem_notes.tns
              python3 file_to_tns.py data.csv --preview
        """)
    )
    ap.add_argument("input",   help="Input file (any supported format)")
    ap.add_argument("output",  nargs="?", help="Output .tns path (default: <input>.tns)")
    ap.add_argument("--preview", action="store_true",
                    help="Print extracted text only — do not write .tns")
    args = ap.parse_args()

    if not os.path.isfile(args.input):
        sys.exit(f"[ERROR] File not found: {args.input}")

    doc_name = os.path.splitext(os.path.basename(args.input))[0]
    out_path = args.output or (doc_name + ".tns")

    print(f"[1/3] Extracting text from: {args.input}")
    try:
        text = extract_text(args.input)
    except Exception as e:
        sys.exit(f"[ERROR] Extraction failed: {e}")

    if not text.strip():
        print("[WARN] No text extracted (scanned PDF or empty file?).")
        text = "(No extractable text found.)"

    print(f"       {len(text):,} characters extracted")

    if args.preview:
        print("\n" + "─" * 60)
        print(text[:4000])
        if len(text) > 4000:
            print(f"\n  … [{len(text)-4000:,} more characters not shown]")
        return

    print(f"[2/3] Building .tns …")
    n = build_tns(text, out_path, doc_name)

    size = os.path.getsize(out_path)
    print(f"[3/3] Written: {out_path}  ({size:,} bytes)")
    print(f"\n✓  Done!  {n} problem(s) packed.")
    print(f'   Transfer "{out_path}" to your TI-Nspire via Student Software.')
    if n > 1:
        print(f"   (File was split into {n} pages due to the ~11.5 KB per-page limit.)")

if __name__ == "__main__":
    main()
